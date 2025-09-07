# services/document_service.py
import re
import pdfplumber
from fastapi import UploadFile
from io import BytesIO
from docx import Document
from typing import List, Dict, Optional
import zipfile
import xml.etree.ElementTree as ET

# Lazy loading of spaCy model
_nlp = None

def get_nlp():
    """Get spaCy model with lazy loading"""
    global _nlp
    if _nlp is None:
        try:
            import spacy
            _nlp = spacy.load("en_core_web_sm")
        except OSError:
            # Fallback if model not available
            _nlp = None
    return _nlp

def extract_text_from_pdf(file: UploadFile) -> str:
    """Extract text from PDF using pdfplumber without temp files"""
    pdf_bytes = BytesIO(file.file.read())
    text = ""
    with pdfplumber.open(pdf_bytes) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text

def extract_text_from_docx(file: UploadFile) -> str:
    """
    Extract text from DOCX using direct XML parsing to handle complex table structures
    """
    docx_bytes = BytesIO(file.file.read())
    
    try:
        # Method 1: Direct XML parsing with better table handling
        with zipfile.ZipFile(docx_bytes, 'r') as zip_file:
            if 'word/document.xml' in zip_file.namelist():
                doc_xml = zip_file.read('word/document.xml')
                root = ET.fromstring(doc_xml)
                
                # Extract text with better structure preservation
                text_parts = []
                
                # Process paragraphs first
                paragraphs = root.findall('.//w:p', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                for para in paragraphs:
                    para_text = ""
                    text_elements = para.findall('.//w:t', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                    if text_elements:
                        para_text = ' '.join([elem.text for elem in text_elements if elem.text])
                        if para_text.strip():
                            text_parts.append(para_text.strip())
                
                # Process tables with better structure
                tables = root.findall('.//w:tbl', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                for table in tables:
                    table_text = []
                    rows = table.findall('.//w:tr', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                    
                    for row in rows:
                        row_text = []
                        cells = row.findall('.//w:tc', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                        
                        for cell in cells:
                            cell_text_elements = cell.findall('.//w:t', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                            cell_text = ' '.join([elem.text for elem in cell_text_elements if elem.text])
                            if cell_text.strip():
                                row_text.append(cell_text.strip())
                        
                        if row_text:
                            table_text.append(' | '.join(row_text))
                    
                    if table_text:
                        text_parts.extend(table_text)
                
                # Combine all text with proper spacing
                all_text = '\n'.join(text_parts)
                
                if len(all_text.strip()) > 100:  # If we got substantial text, use it
                    return all_text
        
        # Method 2: Fallback to python-docx
        docx_bytes.seek(0)  # Reset file pointer
        doc = Document(docx_bytes)
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                text += row_text + "\n"
        
        return text
        
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        # Method 3: Last resort - try to read as bytes and decode
        try:
            docx_bytes.seek(0)
            content = docx_bytes.read()
            # Look for readable text in the content
            text = ""
            for i in range(0, len(content), 2):
                try:
                    char = content[i:i+2].decode('utf-16-le', errors='ignore')
                    if char.isprintable():
                        text += char
                except:
                    continue
            return text
        except:
            raise ValueError(f"Failed to extract text from DOCX file: {str(e)}")

def anonymise_text(text: str) -> str:
    """Automatically anonymize text to remove names and PII while preserving room numbers and case details"""
    if not text:
        return text
    
    anonymized_text = text
    
    # Step 1: Replace email addresses
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    anonymized_text = re.sub(email_pattern, '[EMAIL]', anonymized_text)
    
    # Step 2: Replace phone numbers
    phone_pattern = r'\+?[\d\s\-\(\)]{7,}'
    anonymized_text = re.sub(phone_pattern, '[PHONE]', anonymized_text)
    
    # Step 3: Replace booking references
    booking_pattern = r'\b(?:REF|#|Booking|Reservation|Confirmation)[\s\-]?\d+[A-Za-z0-9]*\b'
    anonymized_text = re.sub(booking_pattern, '[BOOKING_REFERENCE]', anonymized_text)
    
    # Step 4: Replace guest IDs
    guest_id_pattern = r'\b(?:Guest|Customer|Client)\s*ID\s*[#]?\s*(\d+)\b'
    anonymized_text = re.sub(guest_id_pattern, '[GUEST_ID]', anonymized_text)
    
    # Step 5: Use spaCy for name detection and replacement (with fallback)
    try:
        nlp = get_nlp()
        if nlp:
            doc = nlp(anonymized_text)
            # Sort entities by length (longest first) to avoid partial replacements
            entities = sorted(doc.ents, key=lambda x: len(x.text), reverse=True)
            
            for ent in entities:
                if ent.label_ == "PERSON":
                    # Only replace if it looks like a name and hasn't been replaced already
                    if ent.text not in ['[CLIENT_NAME]', '[GUEST_ID]', '[BOOKING_REFERENCE]', '[EMAIL]', '[PHONE]']:
                        anonymized_text = anonymized_text.replace(ent.text, '[CLIENT_NAME]')
    except Exception as e:
        print(f"Warning: spaCy not available for name detection: {e}")
    
    # Step 6: Additional name patterns for comprehensive coverage
    # Look for capitalized words that might be names (but not room numbers or field labels)
    name_patterns = [
        r'\b[A-Z][a-z]{2,}\s+[A-Z][a-z]{2,}\b',  # Two capitalized words
        r'\b(?:Mr\.|Mrs\.|Ms\.|Dr\.|Prof\.)\s+[A-Z][a-z]+\s+[A-Z][a-z]+\b',  # Names with titles
    ]
    
    for pattern in name_patterns:
        potential_names = re.findall(pattern, anonymized_text)
        for name in potential_names:
            # Skip if it's already been replaced or contains common non-name words
            if (name not in ['[CLIENT_NAME]', '[GUEST_ID]', '[BOOKING_REFERENCE]', '[EMAIL]', '[PHONE]'] and
                not any(word in name.lower() for word in ['room', 'floor', 'suite', 'hotel', 'guest', 'check', 'booking', 'maintenance', 'service', 'guest', 'relations', 'report', 'additional', 'notes', 'status', 'type', 'importance', 'source', 'membership', 'action', 'case', 'created', 'modified'])):
                anonymized_text = anonymized_text.replace(name, '[CLIENT_NAME]')
    
    return anonymized_text

def parse_cases(text: str, status_type_info: dict = None) -> list:
    """Convert text into structured list of case dicts - updated for table format"""
    cases = []
    
    # Split text into potential case blocks - improved splitting logic
    # Look for multiple patterns that indicate case boundaries
    case_blocks = re.split(r'(?=Created\s+\d{2}/\d{2}/\d{4})', text)
    
    # If we only got one block, try alternative splitting methods
    if len(case_blocks) <= 1:
        print("DEBUG: Single block detected, trying alternative splitting...")
        # Try splitting by other patterns that might indicate case boundaries
        case_blocks = re.split(r'(?=Guest\s+[A-Z][a-z]+)', text)
        
        if len(case_blocks) <= 1:
            # Try splitting by room numbers
            case_blocks = re.split(r'(?=Room\s+\d+)', text)
            
        if len(case_blocks) <= 1:
            # Try splitting by double newlines or section breaks
            case_blocks = re.split(r'\n\s*\n\s*\n', text)
            
        if len(case_blocks) <= 1:
            # Last resort: split by any date pattern
            case_blocks = re.split(r'(?=\d{2}/\d{2}/\d{4})', text)
    
    print(f"DEBUG: Found {len(case_blocks)} potential case blocks")
    
    case_index = 0
    for i, block in enumerate(case_blocks):
        block = block.strip()
        if not block or len(block) < 50:  # Skip very short blocks
            continue
            
        print(f"DEBUG: Processing block {i+1}, length: {len(block)}")
        print(f"DEBUG: Block content: {block[:200]}...")
        
        # Extract ALL case information using comprehensive patterns
        # Enhanced extraction to capture every detail from the PDF
        
        # Extract guest name - multiple patterns for comprehensive coverage
        guest_match = re.search(r'Guest:\s*([^|]+)', block)
        if not guest_match:
            guest_match = re.search(r'Guest\s+([A-Za-z\s]+)', block)
        if not guest_match:
            guest_match = re.search(r'Guest\s*:\s*([^\n\r]+)', block)
        
        # Extract room number - multiple patterns for comprehensive coverage
        room_match = re.search(r'Room:\s*(\d+)', block)
        if not room_match:
            room_match = re.search(r'Room\s+(\d+)', block)
        if not room_match:
            room_match = re.search(r'Room\s*:\s*(\d+)', block)
        if not room_match:
            # Look for room numbers in various formats
            room_match = re.search(r'(\d{3,4})', block)  # 3-4 digit numbers
        
        # Extract status - comprehensive status extraction
        status_match = re.search(r'Status:\s*(\w+)', block)
        if not status_match:
            status_match = re.search(r'Status\s+(\w+)', block)
        if not status_match:
            status_match = re.search(r'Status\s*:\s*(\w+)', block)
        
        # Extract importance - comprehensive importance extraction
        importance_match = re.search(r'Importance:\s*(\w+)', block)
        if not importance_match:
            importance_match = re.search(r'Importance\s+(\w+)', block)
        if not importance_match:
            importance_match = re.search(r'Importance\s*:\s*(\w+)', block)
        
        # Extract type - comprehensive type extraction
        type_match = re.search(r'Type:\s*(\w+)', block)
        if not type_match:
            type_match = re.search(r'Type\s+(\w+)', block)
        if not type_match:
            type_match = re.search(r'Type\s*:\s*(\w+)', block)
        
        # Extract source - comprehensive source extraction
        source_match = re.search(r'Source:\s*([^|]+)', block)
        if not source_match:
            source_match = re.search(r'Source\s+([^\n\r]+)', block)
        if not source_match:
            source_match = re.search(r'Source\s*:\s*([^\n\r]+)', block)
        
        # Extract membership - comprehensive membership extraction
        membership_match = re.search(r'Member:\s*([^|]+)', block)
        if not membership_match:
            membership_match = re.search(r'Member\s+([^\n\r]+)', block)
        if not membership_match:
            membership_match = re.search(r'Member\s*:\s*([^\n\r]+)', block)
        
        # Extract in/out dates - comprehensive date extraction
        in_out_match = re.search(r'In/Out:\s*([^|]+)', block)
        if not in_out_match:
            in_out_match = re.search(r'In/Out\s+([^\n\r]+)', block)
        
        # Extract modified date - comprehensive date extraction
        modified_match = re.search(r'Updated:\s*([^|]+)', block)
        if not modified_match:
            modified_match = re.search(r'Modified:\s*([^|]+)', block)
        if not modified_match:
            modified_match = re.search(r'Last\s+Updated:\s*([^|]+)', block)
        if not modified_match:
            modified_match = re.search(r'Updated\s+([^\n\r]+)', block)
        
        # Extract created date - comprehensive date extraction
        created_match = re.search(r'Created:\s*([^|]+)', block)
        if not created_match:
            created_match = re.search(r'Date:\s*([^|]+)', block)
        if not created_match:
            created_match = re.search(r'Created\s+([^\n\r]+)', block)
        
        # Extract created by - comprehensive user extraction
        created_by_match = re.search(r'Created\s+by:\s*([^|]+)', block)
        if not created_by_match:
            created_by_match = re.search(r'By:\s*([^|]+)', block)
        if not created_by_match:
            created_by_match = re.search(r'Created\s+by\s+([^\n\r]+)', block)
        
        # Extract modified by - comprehensive user extraction
        modified_by_match = re.search(r'Modified\s+by:\s*([^|]+)', block)
        if not modified_by_match:
            modified_by_match = re.search(r'Updated\s+by:\s*([^|]+)', block)
        if not modified_by_match:
            modified_by_match = re.search(r'Modified\s+by\s+([^\n\r]+)', block)
        if not modified_by_match:
            # Look for staff name patterns at the end of blocks
            staff_match = re.search(r'([A-Za-z\s]+)\s*$', block)
            if staff_match:
                potential_name = staff_match.group(1).strip()
                if len(potential_name) < 50 and re.match(r'^[A-Za-z\s]+$', potential_name):
                    modified_by_match = staff_match
            
        # Extract values from the matches
        guest_value = guest_match.group(1).strip() if guest_match else None
        room_value = room_match.group(1).strip() if room_match else None
        status_value = status_match.group(1).strip() if status_match else None
        importance_value = importance_match.group(1).strip() if importance_match else None
        type_value = type_match.group(1).strip() if type_match else None
        source_value = source_match.group(1).strip() if source_match else None
        membership_value = membership_match.group(1).strip() if membership_match else None
        in_out_value = in_out_match.group(1).strip() if in_out_match else None
        modified_value = modified_match.group(1).strip() if modified_match else None
        modified_by_value = modified_by_match.group(1).strip() if modified_by_match else None
        
        # Extract case description as a separate section
        case_match = None
        
        # Method 1: Look for CASE section
        case_match = re.search(r'CASE\s*\n\s*(.+?)(?=\n\s*ACTION|\n\s*Created|$)', block, re.DOTALL | re.IGNORECASE)
        
        # Method 2: Fallback pattern
        if not case_match:
            case_match = re.search(r'CASE\s+([^A-Z]+?)(?=\s+ACTION|$)', block, re.DOTALL | re.IGNORECASE)
        if not case_match:
            # Look for any substantial text that might be a case description
            # Find the longest text block that's not just field names or action-related
            lines = block.split('\n')
            description_lines = []
            action_keywords = ['update', 'action', 'resolved', 'completed', 'follow-up', 'followup', 'done', 'finished']
            
            for line in lines:
                line = line.strip()
                if (len(line) > 20 and 
                    not re.match(r'^(Created|Guest|Status|Type|Room|Importance|Modified|Source|Membership|IN/OUT|ACTION)', line, re.IGNORECASE) and
                    not re.match(r'^\d{2}/\d{2}/\d{4}', line) and
                    not any(keyword in line.lower() for keyword in action_keywords)):
                    description_lines.append(line)
            
            if description_lines:
                case_description = ' '.join(description_lines[:3])  # Take first 3 meaningful lines
                if len(case_description) > 10:
                    case_match = type('MockMatch', (), {'group': lambda self, x: case_description})()
        
        # Extract action as a separate section - simplified approach
        action_match = None
        
        # Method 1: Look for ACTION section
        action_match = re.search(r'ACTION\s*\n\s*(.+?)(?=\n\s*Created|\n\s*CASE|$)', block, re.DOTALL | re.IGNORECASE)
        
        # Method 2: Look for Update patterns
        if not action_match:
            action_match = re.search(r'Update:\s*(.+?)(?=\n\s*Created|\n\s*CASE|$)', block, re.DOTALL)
        
        # Method 3: Look for Action taken/required
        if not action_match:
            action_match = re.search(r'Action\s+(?:taken|required):\s*(.+?)(?=\n\s*Created|\n\s*CASE|$)', block, re.DOTALL | re.IGNORECASE)
        
        # Method 4: Look for lines with action keywords
        if not action_match:
            lines = block.split('\n')
            action_lines = []
            action_keywords = ['update', 'action', 'resolved', 'completed', 'follow-up', 'followup', 'done', 'finished']
            
            for line in lines:
                line = line.strip()
                if (len(line) > 15 and 
                    any(keyword in line.lower() for keyword in action_keywords) and
                    not re.match(r'^(Created|Guest|Status|Type|Room|Importance|Modified|Source|Membership|IN/OUT|CASE)', line, re.IGNORECASE) and
                    not re.match(r'^\d{2}/\d{2}/\d{4}', line)):
                    action_lines.append(line)
            
            if action_lines:
                action_description = ' '.join(action_lines[:2])
                if len(action_description) > 10:
                    action_match = type('MockMatch', (), {'group': lambda self, x: action_description})()
        
        # Also look for IN/OUT dates
        in_out_match = re.search(r'IN/OUT\s*\n\s*([^\n\r]+)', block)
        if not in_out_match:
            in_out_match = re.search(r'IN/OUT\s+([^\n\r]+)', block)
        
        # Create case object with guaranteed title and more fields

        # Clean up guest name - remove pipe-separated values that indicate malformed data
        if guest_value and '|' in guest_value:
            # Take only the first part before the pipe
            guest_value = guest_value.split('|')[0].strip()
        
        # Use room number as title (primary requirement)
        if room_value:
            title = f"Room {room_value}"
        elif case_match and case_match.group(1).strip():
            # Use first 50 characters of case description as fallback
            desc = case_match.group(1).strip()
            title = desc[:50] + "..." if len(desc) > 50 else desc
        else:
            title = "Untitled Case"
        
        case = {
            "created": created_match.group(1).strip() if created_match else None,
            "status": status_value,
            "created_by": created_by_match.group(1).strip() if created_by_match else None,
            "room": room_value,
            "importance": importance_value,
            "modified": modified_match.group(1).strip() if modified_match else None,
            "modified_by": modified_by_match.group(1).strip() if modified_by_match else None,
            "source": source_value,
            "membership": membership_value,
            "type": type_value,
            "case_description": case_match.group(1).strip() if case_match else None,
            "action": action_match.group(1).strip() if action_match else None,
            "in_out": in_out_value,
            "title": title  # Always guaranteed to have a value (room number)
        }
        
        print(f"DEBUG: Extracted case: {case}")
        if action_match:
            print(f"DEBUG: Found action: {action_match.group(1).strip()[:100]}...")
        if case_match:
            print(f"DEBUG: Found case description: {case_match.group(1).strip()[:100]}...")
        
        # More lenient validation - accept cases with any meaningful information
        # Check if we have at least one of: room number, case description, or title
        has_meaningful_data = (
            (case["room"] and str(case["room"]).strip()) or
            (case["case_description"] and len(case["case_description"]) > 10) or
            (case["title"] and case["title"] != "Untitled Case" and len(case["title"]) > 5)
        )
        
        if has_meaningful_data:
            cases.append(case)
            case_index += 1  # Increment case index for status/type mapping
            print(f"DEBUG: Added case {len(cases)}")
        else:
            print(f"DEBUG: Skipped case - insufficient information or malformed data")
    
    print(f"DEBUG: Total cases found: {len(cases)}")
    return cases

def process_document(file: UploadFile) -> list:
    """Optimized pipeline: extract → try AI parsing → fallback to regex parsing"""
    try:
        # Extract text once and cache it
        if file.filename.lower().endswith('.pdf'):
            raw_text = extract_text_from_pdf(file)
        elif file.filename.lower().endswith('.docx'):
            raw_text = extract_text_from_docx(file)
        elif file.filename.lower().endswith('.txt'):
            # For text files, just read the content directly
            try:
                raw_text = file.file.read().decode('utf-8')
            except UnicodeDecodeError:
                # Try different encodings if utf-8 fails
                file.file.seek(0)  # Reset file pointer
                try:
                    raw_text = file.file.read().decode('latin-1')
                except:
                    file.file.seek(0)
                    raw_text = file.file.read().decode('cp1252')
        else:
            raise ValueError(f"Unsupported file type: {file.filename}")
        
        # Quick validation
        if not raw_text or len(raw_text.strip()) < 10:
            return []
        
        # Apply automatic anonymization to protect privacy
        anonymized_text = anonymise_text(raw_text)
        
        # Skip AI parsing - use regex parsing directly
        # AI parsing disabled per user preference
        
        # Fallback to optimized regex parsing with anonymized text
        status_type_info = extract_status_type_info(anonymized_text)
        cases = parse_cases(anonymized_text, status_type_info)
        
        return cases
        
    except Exception as e:
        print(f"ERROR in process_document: {e}")
        raise

def extract_status_type_info(text: str) -> dict:
    """Extract status and type information from original text before anonymization"""
    info = {}
    
    # Find all Status and Type fields in the text
    status_matches = re.findall(r'Status\s*\n\s*(\w+)', text)
    type_matches = re.findall(r'Type\s*\n\s*(\w+)', text)
    
    # Create a mapping of positions to values
    for i, status in enumerate(status_matches):
        info[f'status_{i}'] = status
    
    for i, type_val in enumerate(type_matches):
        info[f'type_{i}'] = type_val
    
    print(f"DEBUG: Extracted status/type info: {info}")
    return info

# Legacy functions for backward compatibility
def process_pdf(file: UploadFile) -> list:
    """Legacy function - now redirects to process_document"""
    return process_document(file)

def extract_text_from_document(file: UploadFile) -> str:
    """Extract text from either PDF or DOCX file based on file extension"""
    if file.filename.lower().endswith('.pdf'):
        return extract_text_from_pdf(file)
    elif file.filename.lower().endswith('.docx'):
        return extract_text_from_docx(file)
    elif file.filename.lower().endswith('.txt'):
        try:
            return file.file.read().decode('utf-8')
        except UnicodeDecodeError:
            # Try different encodings if utf-8 fails
            file.file.seek(0)  # Reset file pointer
            try:
                return file.file.read().decode('latin-1')
            except:
                file.file.seek(0)
                return file.file.read().decode('cp1252')
    else:
        raise ValueError(f"Unsupported file type: {file.filename}")