# services/case_parser_service.py
import re
from typing import List, Dict, Any, Optional
from docx import Document
import pdfplumber
from fastapi import UploadFile
from io import BytesIO

class CaseParserService:
    """Enhanced case parsing service for various document formats"""
    
    def __init__(self):
        # Common case patterns to look for
        self.case_patterns = {
            'room_number': [
                r'\b(?:Room|Rm|Room\s*#?)\s*(\d+)\b',
                r'\b(\d+)\s*(?:Room|Rm)\b',
                r'\bRoom\s*(\d+)\b'
            ],
            'guest_name': [
                r'\b(?:Guest|Customer|Client)\s*:\s*([^\n\r,]+)',
                r'\b([A-Z][a-z]+\s+[A-Z][a-z]+)\s*(?:checked\s+in|arrived|staying)',
                r'\b(?:Mr\.|Mrs\.|Ms\.|Dr\.)\s*([A-Z][a-z]+\s+[A-Z][a-z]+)'
            ],
            'status': [
                r'\b(?:Status|State)\s*:\s*(\w+)',
                r'\b(?:Open|Closed|Pending|Resolved|Active|Inactive)\b',
                r'\b(?:Case\s+)?Status\s*:\s*(\w+)'
            ],
            'importance': [
                r'\b(?:Importance|Priority|Urgency)\s*:\s*(\w+)',
                r'\b(?:High|Medium|Low|Critical|Urgent)\b',
                r'\b(?:Priority\s+)?Level\s*:\s*(\w+)'
            ],
            'case_type': [
                r'\b(?:Type|Category)\s*:\s*(\w+)',
                r'\b(?:General|Complaint|Request|Issue|Maintenance|Service)\b',
                r'\b(?:Case\s+)?Type\s*:\s*(\w+)'
            ],
            'description': [
                r'\b(?:Description|Details|Summary)\s*:\s*([^\n\r]+)',
                r'\b(?:Issue|Problem|Request)\s*:\s*([^\n\r]+)',
                r'\b(?:Case|Matter)\s*:\s*([^\n\r]+)'
            ],
            'action': [
                r'\b(?:Action|Next\s+Steps|Resolution)\s*:\s*([^\n\r]+)',
                r'\b(?:To\s+Do|Follow\s+Up)\s*:\s*([^\n\r]+)',
                r'\b(?:Plan|Solution)\s*:\s*([^\n\r]+)'
            ]
        }
    
    def parse_document_for_cases(self, file: UploadFile) -> Dict[str, Any]:
        """
        Parse document and extract case information
        
        Returns:
            Dictionary with parsed cases and metadata
        """
        try:
            if file.filename.lower().endswith('.docx'):
                return self._parse_docx_for_cases(file)
            elif file.filename.lower().endswith('.pdf'):
                return self._parse_pdf_for_cases(file)
            else:
                raise ValueError(f"Unsupported file type: {file.filename}")
        except Exception as e:
            raise Exception(f"Error parsing document: {str(e)}")
    
    def _parse_docx_for_cases(self, file: UploadFile) -> Dict[str, Any]:
        """Parse DOCX document for case information"""
        docx_bytes = BytesIO(file.file.read())
        doc = Document(docx_bytes)
        
        # Extract all text
        all_text = ""
        for para in doc.paragraphs:
            all_text += para.text + "\n"
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text + " | "
                all_text += "\n"
        
        return self._extract_cases_from_text(all_text, file.filename)
    
    def _parse_pdf_for_cases(self, file: UploadFile) -> Dict[str, Any]:
        """Parse PDF document for case information"""
        pdf_bytes = BytesIO(file.file.read())
        
        all_text = ""
        with pdfplumber.open(pdf_bytes) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    all_text += page_text + "\n"
        
        return self._extract_cases_from_text(all_text, file.filename)
    
    def _extract_cases_from_text(self, text: str, filename: str) -> Dict[str, Any]:
        """Extract case information from text using pattern matching"""
        
        # Check if text has sufficient content
        if len(text.strip()) < 50:
            return {
                "filename": filename,
                "cases_found": 0,
                "cases": [],
                "warnings": ["Document appears to have insufficient content for case parsing"],
                "suggestions": [
                    "Document may be empty or contain only headers",
                    "Consider manually inputting case information",
                    "Check if document format is supported"
                ],
                "raw_text": text.strip(),
                "text_length": len(text.strip())
            }
        
        # Try to find case information using patterns
        cases = []
        warnings = []
        
        # Look for structured case data
        structured_cases = self._find_structured_cases(text)
        if structured_cases:
            cases.extend(structured_cases)
        
        # Look for unstructured case data
        unstructured_cases = self._find_unstructured_cases(text)
        if unstructured_cases:
            cases.extend(unstructured_cases)
        
        # If no cases found, try to create a default case from available information
        if not cases:
            default_case = self._create_default_case_from_text(text)
            if default_case:
                cases.append(default_case)
                warnings.append("No structured case data found, created default case from available information")
        
        # Analyze text for potential PII
        pii_analysis = self._analyze_text_for_pii(text)
        
        return {
            "filename": filename,
            "cases_found": len(cases),
            "cases": cases,
            "warnings": warnings,
            "suggestions": self._generate_suggestions(text, cases),
            "raw_text": text[:500] + "..." if len(text) > 500 else text,
            "text_length": len(text),
            "pii_analysis": pii_analysis
        }
    
    def _find_structured_cases(self, text: str) -> List[Dict[str, Any]]:
        """Find cases with structured data"""
        cases = []
        
        # Split text into potential case blocks
        # Look for common separators
        separators = [
            r'\n\s*\n',  # Double newlines
            r'\n\s*[-*•]\s*',  # Bullet points
            r'\n\s*\d+\.\s*',  # Numbered lists
            r'\n\s*Case\s*\d*\s*:',  # Case headers
            r'\n\s*Guest\s*:',  # Guest headers
        ]
        
        # Try different splitting strategies
        for separator in separators:
            blocks = re.split(separator, text)
            if len(blocks) > 1:
                for block in blocks:
                    if len(block.strip()) > 20:  # Minimum block size
                        case = self._parse_case_block(block)
                        if case:
                            cases.append(case)
                if cases:
                    break
        
        return cases
    
    def _find_unstructured_cases(self, text: str) -> List[Dict[str, Any]]:
        """Find cases in unstructured text"""
        cases = []
        
        # Look for any text that might contain case information
        lines = text.split('\n')
        current_case = {}
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if line contains case information
            for field, patterns in self.case_patterns.items():
                for pattern in patterns:
                    match = re.search(pattern, line, re.IGNORECASE)
                    if match:
                        if field == 'description' or field == 'action':
                            current_case[field] = match.group(1).strip()
                        else:
                            current_case[field] = match.group(1).strip()
                        break
            
            # If we have enough information, create a case
            if len(current_case) >= 2:  # At least 2 fields
                case = self._create_case_from_dict(current_case)
                if case:
                    cases.append(case)
                    current_case = {}
        
        return cases
    
    def _parse_case_block(self, block: str) -> Optional[Dict[str, Any]]:
        """Parse a single case block"""
        case_data = {}
        
        # Extract information using patterns
        for field, patterns in self.case_patterns.items():
            for pattern in patterns:
                match = re.search(pattern, block, re.IGNORECASE)
                if match:
                    if field == 'description' or field == 'action':
                        case_data[field] = match.group(1).strip()
                    else:
                        case_data[field] = match.group(1).strip()
                    break
        
        return self._create_case_from_dict(case_data) if case_data else None
    
    def _create_case_from_dict(self, case_data: Dict[str, Any]) -> Dict[str, Any]:
        """Create a standardized case from extracted data"""
        case = {
            "room": case_data.get('room_number'),
            "status": case_data.get('status', 'Open'),
            "importance": case_data.get('importance', 'Medium'),
            "type": case_data.get('case_type', 'General'),
            "title": case_data.get('description', 'Untitled Case'),
            "action": case_data.get('action'),
            "guest_name": case_data.get('guest_name')
        }
        
        # Ensure title is never empty
        if not case['title'] or case['title'].strip() == '':
            case['title'] = "Untitled Case"
        
        return case
    
    def _create_default_case_from_text(self, text: str) -> Optional[Dict[str, Any]]:
        """Create a default case when no structured data is found"""
        # Look for any meaningful information in the text
        lines = text.split('\n')
        meaningful_lines = [line.strip() for line in lines if len(line.strip()) > 10]
        
        if meaningful_lines:
            # Use the first meaningful line as description
            description = meaningful_lines[0][:100] + "..." if len(meaningful_lines[0]) > 100 else meaningful_lines[0]
            
            return {
                "room": None,
                "status": "Open",
                "importance": "Medium",
                "type": "General",
                "title": description,
                "action": None,
                "guest_name": None
            }
        
        return None
    
    def _analyze_text_for_pii(self, text: str) -> Dict[str, Any]:
        """Analyze text for potential PII"""
        pii_types = {
            'emails': len(re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)),
            'phone_numbers': len(re.findall(r'\+?[\d\s\-\(\)]{7,}', text)),
            'names': len(re.findall(r'\b(?:Mr\.|Mrs\.|Ms\.|Dr\.)\s*[A-Z][a-z]+\s+[A-Z][a-z]+', text)),
            'room_numbers': len(re.findall(r'\b(?:Room|Rm)\s*\d+', text)),
            'dates': len(re.findall(r'\b\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}\b', text))
        }
        
        total_pii = sum(pii_types.values())
        
        return {
            "total_pii_found": total_pii,
            "pii_breakdown": pii_types,
            "has_pii": total_pii > 0
        }
    
    def _generate_suggestions(self, text: str, cases: List[Dict[str, Any]]) -> List[str]:
        """Generate suggestions based on document content"""
        suggestions = []
        
        if not cases:
            suggestions.append("No cases found - consider manually inputting case information")
            suggestions.append("Check if document contains case data in expected format")
        
        if len(text.strip()) < 100:
            suggestions.append("Document appears to be minimal - may need additional content")
        
        if cases:
            suggestions.append(f"Successfully parsed {len(cases)} cases")
            suggestions.append("Review extracted information for accuracy")
        
        return suggestions

# Create a global instance
case_parser_service = CaseParserService()
