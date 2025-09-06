# services/anonymization_service.py
import re
from typing import Dict, List, Any, Optional
from fastapi import UploadFile
from io import BytesIO
from docx import Document
import pdfplumber

# Lazy loading of spaCy model
_nlp = None

def get_nlp():
    """Get spaCy model with lazy loading"""
    global _nlp
    if _nlp is None:
        try:
            import spacy
            _nlp = spacy.load("en_core_web_sm")
        except OSError:
            # Fallback if model not available
            _nlp = None
    return _nlp

class AnonymizationService:
    """Enhanced anonymization service for documents with focus on client names and room information"""
    
    def __init__(self):
        # Enhanced patterns for PII detection with focus on hotel/guest relations
        self.patterns = {
            # Personal Information
            'email': r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
            'phone': r'\+?[\d\s\-\(\)]{7,}',
            'credit_card': r'\b\d{4}[\s\-]?\d{4}[\s\-]?\d{4}[\s\-]?\d{4}\b',
            'passport': r'\b[A-Z]{1,2}\d{6,9}\b',
            'ssn': r'\b\d{3}[\-]?\d{2}[\-]?\d{4}\b',
            
            # Hotel-specific patterns
            'booking_ref': r'\b(?:REF|#|Booking|Reservation|Confirmation)[\s\-]?\d+[A-Za-z0-9]*\b',
            'guest_id': r'\b(?:Guest|Customer|Client)\s*ID\s*[#]?\s*(\d+)\b',
            'reservation_id': r'\b(?:Reservation|Booking)\s*ID\s*[#]?\s*(\d+)\b',
            
            # Room information - Enhanced patterns
            'room_number': r'\b(?:Room|Rm|Suite|Apartment|Apt)\s*[#]?\s*(\d+[A-Za-z]?)\b',
            'room_type': r'\b(?:Standard|Deluxe|Suite|Executive|Presidential|King|Queen|Twin|Single|Double)\s+(?:Room|Suite|Apartment)\b',
            'floor_number': r'\b(?:Floor|Level|Story)\s*(\d+)\b',
            'building_section': r'\b(?:Building|Tower|Wing|Section)\s*[A-Za-z]\b',
            
            # Check-in/Check-out information
            'check_in': r'\b(?:Check[-\s]?in|Arrival|Check[-\s]?in\s*Date)\s*[:=]?\s*([^\n]+)\b',
            'check_out': r'\b(?:Check[-\s]?out|Departure|Check[-\s]?out\s*Date)\s*[:=]?\s*([^\n]+)\b',
            
            # Dates and Times
            'date': r'\b\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}\b',
            'time': r'\b\d{1,2}:\d{2}(?::\d{2})?\s*(?:AM|PM|am|pm)?\b',
            
            # Address and Location
            'address': r'\b\d+\s+[A-Za-z\s]+(?:Street|St|Avenue|Ave|Road|Rd|Boulevard|Blvd|Lane|Ln|Drive|Dr)\b',
            'postal_code': r'\b[A-Z]{1,2}\d[A-Z]\s?\d[A-Z]\d\b',  # UK format
            'ip_address': r'\b\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}\b',
            'url': r'https?://(?:[-\w.])+(?:[:\d]+)?(?:/(?:[\w/_.])*(?:\?(?:[\w&=%.])*)?(?:#(?:[\w.])*)?)?',
            
            # Additional hotel-specific patterns
            'amenity_request': r'\b(?:Request|Need|Would like)\s+(?:extra|additional|more)\s+(?:towels|pillows|blankets|toiletries)\b',
            'maintenance_request': r'\b(?:Issue|Problem|Broken|Not working|Faulty)\s+(?:with|in)\s+(?:AC|heating|plumbing|electrical|appliance)\b',
            'service_request': r'\b(?:Room service|Housekeeping|Concierge|Maintenance|Technical support)\s+(?:request|call|assistance)\b',
        }
        
        # Enhanced replacement tokens
        self.replacements = {
            'email': '[EMAIL]',
            'phone': '[PHONE]',
            'booking_ref': '[BOOKING_REFERENCE]',
            'credit_card': '[CREDIT_CARD]',
            'passport': '[PASSPORT]',
            'ssn': '[SSN]',
            'guest_id': '[GUEST_ID]',
            'reservation_id': '[RESERVATION_ID]',
            'date': '[DATE]',
            'time': '[TIME]',
            'address': '[ADDRESS]',
            'postal_code': '[POSTAL_CODE]',
            'ip_address': '[IP_ADDRESS]',
            'url': '[URL]',
            'name': '[CLIENT_NAME]',
            'company': '[COMPANY]',
            'location': '[LOCATION]',
            'room_number': '[ROOM_NUMBER]',
            'room_type': '[ROOM_TYPE]',
            'floor_number': '[FLOOR_NUMBER]',
            'building_section': '[BUILDING_SECTION]',
            'check_in': '[CHECK_IN_DATE]',
            'check_out': '[CHECK_OUT_DATE]',
            'amenity_request': '[AMENITY_REQUEST]',
            'maintenance_request': '[MAINTENANCE_REQUEST]',
            'service_request': '[SERVICE_REQUEST]',
        }
    
    def anonymize_text(self, text: str, preserve_dates: bool = False, preserve_times: bool = False) -> str:
        """
        Anonymize text content with GDPR compliance - remove names but preserve room numbers
        
        Args:
            text: Input text to anonymize
            preserve_dates: Whether to preserve date information
            preserve_times: Whether to preserve time information
        
        Returns:
            Anonymized text with names removed but room numbers preserved
        """
        if not text:
            return text
            
        anonymized_text = text
        
        # Step 1: Replace names with titles (Mr., Mrs., etc.) - GDPR compliance
        name_with_title_pattern = r'\b(?:Mr\.|Mrs\.|Ms\.|Dr\.|Prof\.)\s+[A-Z][a-z]+\s+[A-Z][a-z]+\b'
        anonymized_text = re.sub(name_with_title_pattern, '[CLIENT_NAME]', anonymized_text)
        
        # Step 2: Replace guest IDs - GDPR compliance
        guest_id_pattern = r'\b(?:Guest|Customer|Client)\s*ID\s*[#]?\s*(\d+)\b'
        anonymized_text = re.sub(guest_id_pattern, '[GUEST_ID]', anonymized_text)
        
        # Step 3: Replace booking references - GDPR compliance
        booking_pattern = r'\b(?:REF|#|Booking|Reservation|Confirmation)[\s\-]?\d+[A-Za-z0-9]*\b'
        anonymized_text = re.sub(booking_pattern, '[BOOKING_REFERENCE]', anonymized_text)
        
        # Step 4: Replace email addresses - GDPR compliance
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        anonymized_text = re.sub(email_pattern, '[EMAIL]', anonymized_text)
        
        # Step 5: Replace phone numbers - GDPR compliance
        phone_pattern = r'\+?[\d\s\-\(\)]{7,}'
        anonymized_text = re.sub(phone_pattern, '[PHONE]', anonymized_text)
        
        # Step 6: Replace dates and times (if not preserved)
        if not preserve_dates:
            date_pattern = r'\b\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}\b'
            anonymized_text = re.sub(date_pattern, '[DATE]', anonymized_text)
        
        if not preserve_times:
            time_pattern = r'\b\d{1,2}:\d{2}(?::\d{2})?\s*(?:AM|PM|am|pm)?\b'
            anonymized_text = re.sub(time_pattern, '[TIME]', anonymized_text)
        
        # Step 7: Use spaCy for additional name detection (fallback) - GDPR compliance
        nlp = get_nlp()
        if nlp:
            doc = nlp(anonymized_text)
            # Sort entities by length (longest first) to avoid partial replacements
            entities = sorted(doc.ents, key=lambda x: len(x.text), reverse=True)
            
            for ent in entities:
                if ent.label_ == "PERSON":
                    # Only replace if it looks like a name and hasn't been replaced already
                    if ent.text not in ['[CLIENT_NAME]', '[GUEST_ID]', '[BOOKING_REFERENCE]']:
                        anonymized_text = anonymized_text.replace(ent.text, '[CLIENT_NAME]')
        
        # Step 8: Additional name patterns for GDPR compliance
        # Look for capitalized words that might be names (but not room numbers)
        name_patterns = [
            r'\b[A-Z][a-z]{2,}\s+[A-Z][a-z]{2,}\b',  # Two capitalized words
        ]
        
        for pattern in name_patterns:
            potential_names = re.findall(pattern, anonymized_text)
            for name in potential_names:
                # Skip if it's already been replaced or contains common non-name words
                if (name not in ['[CLIENT_NAME]', '[GUEST_ID]', '[BOOKING_REFERENCE]', '[EMAIL]', '[PHONE]'] and
                    not any(word in name.lower() for word in ['room', 'floor', 'suite', 'hotel', 'guest', 'check', 'booking', 'maintenance', 'service', 'guest', 'relations', 'report', 'additional', 'notes'])):
                    anonymized_text = anonymized_text.replace(name, '[CLIENT_NAME]')
        
        return anonymized_text
    
    def anonymize_document(self, file: UploadFile, preserve_dates: bool = False, preserve_times: bool = False) -> Dict[str, Any]:
        """
        Anonymize an entire document (PDF or DOCX)
        
        Args:
            file: Uploaded file to anonymize
            preserve_dates: Whether to preserve date information
            preserve_times: Whether to preserve time information
        
        Returns:
            Dictionary containing original and anonymized content
        """
        try:
            if file.filename.lower().endswith('.docx'):
                return self._anonymize_docx(file, preserve_dates, preserve_times)
            elif file.filename.lower().endswith('.pdf'):
                return self._anonymize_pdf(file, preserve_dates, preserve_times)
            else:
                raise ValueError(f"Unsupported file type: {file.filename}")
        except Exception as e:
            raise Exception(f"Error anonymizing document: {str(e)}")
    
    def _anonymize_docx(self, file: UploadFile, preserve_dates: bool, preserve_times: bool) -> Dict[str, Any]:
        """Anonymize DOCX document"""
        docx_bytes = BytesIO(file.file.read())
        doc = Document(docx_bytes)
        
        original_content = []
        anonymized_content = []
        
        # Process paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                original_content.append(paragraph.text)
                anonymized_text = self.anonymize_text(paragraph.text, preserve_dates, preserve_times)
                anonymized_content.append(anonymized_text)
        
        # Process tables
        for table in doc.tables:
            table_data = []
            anonymized_table_data = []
            
            for row in table.rows:
                row_data = []
                anonymized_row_data = []
                
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                    anonymized_cell_text = self.anonymize_text(cell_text, preserve_dates, preserve_times)
                    anonymized_row_data.append(anonymized_cell_text)
                
                table_data.append(row_data)
                anonymized_table_data.append(anonymized_row_data)
            
            original_content.append(f"TABLE: {table_data}")
            anonymized_content.append(f"TABLE: {anonymized_table_data}")
        
        return {
            "filename": file.filename,
            "file_type": "docx",
            "original_content": original_content,
            "anonymized_content": anonymized_content,
            "anonymization_summary": self._generate_summary(original_content, anonymized_content)
        }
    
    def _anonymize_pdf(self, file: UploadFile, preserve_dates: bool, preserve_times: bool) -> Dict[str, Any]:
        """Anonymize PDF document"""
        pdf_bytes = BytesIO(file.file.read())
        
        original_content = []
        anonymized_content = []
        
        with pdfplumber.open(pdf_bytes) as pdf:
            for page_num, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    original_content.append(f"Page {page_num + 1}: {page_text}")
                    anonymized_text = self.anonymize_text(page_text, preserve_dates, preserve_times)
                    anonymized_content.append(f"Page {page_num + 1}: {anonymized_text}")
        
        return {
            "filename": file.filename,
            "file_type": "pdf",
            "original_content": original_content,
            "anonymized_content": anonymized_content,
            "anonymization_summary": self._generate_summary(original_content, anonymized_content)
        }
    
    def _generate_summary(self, original_content: List[str], anonymized_content: List[str]) -> Dict[str, Any]:
        """Generate anonymization summary"""
        original_text = " ".join(original_content)
        anonymized_text = " ".join(anonymized_content)
        
        # Count replacements
        replacements = {}
        for pattern_name, replacement in self.replacements.items():
            count = anonymized_text.count(replacement)
            if count > 0:
                replacements[pattern_name] = count
        
        return {
            "total_replacements": sum(replacements.values()),
            "replacement_breakdown": replacements,
            "anonymization_ratio": len([c for c in anonymized_text if c == '[']) / len(anonymized_text) if anonymized_text else 0
        }
    
    def get_anonymization_stats(self, text: str) -> Dict[str, Any]:
        """Get statistics about what would be anonymized in a text"""
        if not text:
            return {"total_potential_pii": 0, "breakdown": {}}
        
        stats = {}
        total_count = 0
        
        for pattern_name, pattern in self.patterns.items():
            matches = re.findall(pattern, text, flags=re.IGNORECASE)
            count = len(matches)
            if count > 0:
                stats[pattern_name] = {
                    "count": count,
                    "examples": matches[:3]  # Show first 3 examples
                }
                total_count += count
        
        return {
            "total_potential_pii": total_count,
            "breakdown": stats
        }

# Create a global instance
anonymization_service = AnonymizationService()
